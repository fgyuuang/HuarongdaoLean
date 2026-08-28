# -*- coding: utf-8 -*-
"""Build the maintained HuarongdaoLean version-audit DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "VERSION_AUDIT.md"
OUTPUT = ROOT / "docs" / "HuarongdaoLean-Version-Audit-2026-08-28.docx"
ASSETS = ROOT / "docs" / "assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172A3A"
MUTED = "5F6B75"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F3EC"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
BLACK = "000000"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def font_path(bold: bool = False) -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    raise FileNotFoundError("No suitable Windows font found")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold), size=size)


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int],
                fill: str, outline: str, radius: int = 16, width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}",
                           outline=f"#{outline}", width=width)


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int],
                  title: str, subtitle: str, fill: str = INK) -> None:
    title_font = load_font(31, True)
    subtitle_font = load_font(21)
    left, top, right, bottom = box
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    subtitle_bbox = draw.multiline_textbbox(
        (0, 0), subtitle, font=subtitle_font, spacing=5, align="center")
    total_height = (title_bbox[3] - title_bbox[1]) + 12 + (
        subtitle_bbox[3] - subtitle_bbox[1])
    y = top + (bottom - top - total_height) / 2
    draw.text(((left + right) / 2, y), title, font=title_font,
              fill=f"#{fill}", anchor="ma")
    draw.multiline_text(((left + right) / 2, y + 52), subtitle,
                        font=subtitle_font, fill=f"#{MUTED}",
                        anchor="ma", align="center", spacing=5)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int],
          end: tuple[int, int], color: str = BLUE, width: int = 5) -> None:
    draw.line([start, end], fill=f"#{color}", width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    length = max(1.0, (dx * dx + dy * dy) ** 0.5)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 15
    points = [
        (ex, ey),
        (ex - ux * size + px * size * 0.55,
         ey - uy * size + py * size * 0.55),
        (ex - ux * size - px * size * 0.55,
         ey - uy * size - py * size * 0.55),
    ]
    draw.polygon(points, fill=f"#{color}")


def build_lean_architecture(path: Path) -> None:
    image = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, True)
    draw.text((80, 55), "Lean 4 形式化与有限表示依赖图",
              font=title_font, fill=f"#{INK}")

    left_boxes = [
        ((90, 155, 555, 325), "经典规则层",
         "Model / Transition / Paths\ntryMove · ValidState · Path"),
        ((90, 405, 555, 575), "通用规则层",
         "Generic Model / Paths\nPuzzleSpec · goalMatches"),
    ]
    center_boxes = [
        ((665, 150, 1160, 335), "StateSpace.Task",
         "initial · goal · step\nWalk · Reachable · Solution"),
        ((665, 400, 1160, 585), "有限表示与证书",
         "StateGraph → Fin SimpleGraph\nclosure · BFS dist · cuts"),
        ((665, 650, 1160, 835), "商与压缩",
         "shape → mirror → corridor\nBisimulation · weighted lift"),
    ]
    right_boxes = [
        ((1270, 155, 1710, 325), "Mathlib 图论",
         "SimpleGraph.Metric\ndist · Reachable · shortest walk"),
        ((1270, 405, 1710, 575), "Mathlib 群作用",
         "MulAction.orbitRel\norbit · stabilizer · quotient"),
        ((1270, 655, 1710, 825), "可执行入口",
         "ExportMain / CertMain\nGenericMain / JSON"),
    ]
    for box, title, subtitle in left_boxes:
        rounded_box(draw, box, LIGHT_GRAY, "AAB4BD")
        centered_text(draw, box, title, subtitle)
    for box, title, subtitle in center_boxes:
        rounded_box(draw, box, LIGHT_BLUE, BLUE)
        centered_text(draw, box, title, subtitle)
    for box, title, subtitle in right_boxes:
        rounded_box(draw, box, PALE_GREEN, "5A8E69")
        centered_text(draw, box, title, subtitle)

    arrow(draw, (555, 240), (665, 240))
    arrow(draw, (555, 490), (665, 490))
    arrow(draw, (912, 335), (912, 400))
    arrow(draw, (912, 585), (912, 650))
    arrow(draw, (1160, 240), (1270, 240))
    arrow(draw, (1160, 490), (1270, 490))
    arrow(draw, (1160, 742), (1270, 742))
    arrow(draw, (1490, 575), (1490, 655))

    note_font = load_font(22)
    draw.text((90, 930),
              "核心原则：Task 是数学对象；Graph/StateGraph 是有限表示；JSON 与坐标是运行时产物。",
              font=note_font, fill=f"#{DARK_BLUE}")
    draw.text((90, 980),
              "证明链只在明确的投影、双模拟、checker soundness 或路径提升定理处跨层。",
              font=note_font, fill=f"#{MUTED}")
    image.save(path, dpi=(160, 160))


def build_runtime_pipeline(path: Path) -> None:
    image = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, True)
    draw.text((80, 55), "本地生成、布局与交互链",
              font=title_font, fill=f"#{INK}")

    boxes = [
        ((80, 190, 355, 390), "浏览器规格",
         "PuzzleSpec 编辑\n经典层选择"),
        ((430, 190, 705, 390), "本地 Node",
         "参数校验\nHTTP / 静态文件"),
        ((780, 190, 1055, 390), "Lean 可执行",
         "BFS / A*\nchecker / JSON"),
        ((1130, 190, 1405, 390), "布局 Worker",
         "地标距离\n4D 力 → 3D"),
        ((1480, 190, 1755, 390), "WebGL 渲染",
         "Three.js\n3d-force-graph"),
    ]
    fills = [LIGHT_GRAY, LIGHT_GRAY, LIGHT_BLUE, PALE_GOLD, PALE_GREEN]
    outlines = ["AAB4BD", "AAB4BD", BLUE, "B78A28", "5A8E69"]
    for (box, title, subtitle), fill, outline in zip(boxes, fills, outlines):
        rounded_box(draw, box, fill, outline)
        centered_text(draw, box, title, subtitle)
    for first, second in zip(boxes, boxes[1:]):
        arrow(draw, (first[0][2], 290), (second[0][0], 290))

    lower = [
        ((160, 585, 570, 770), "形式化数据",
         "状态 · 动作 · 边\n目标 · 距离 · 宏边步骤"),
        ((695, 585, 1105, 770), "纯可视化数据",
         "x/y/z · 相机 · 颜色\n动画 · 屏幕命中"),
        ((1230, 585, 1640, 770), "交互回路",
         "点击选端点 → 图上 BFS\n逐条已有有向边播放"),
    ]
    for box, title, subtitle in lower:
        rounded_box(draw, box, WHITE, BLUE)
        centered_text(draw, box, title, subtitle)
    arrow(draw, (917, 390), (365, 585), DARK_BLUE)
    arrow(draw, (1267, 390), (900, 585), "B78A28")
    arrow(draw, (1617, 390), (1435, 585), "5A8E69")
    arrow(draw, (1230, 677), (1105, 677), BLUE)
    arrow(draw, (695, 677), (570, 677), BLUE)

    note_font = load_font(22)
    draw.text((115, 885),
              "坐标可以被释放、重新加热和恢复；这些操作不改变 Lean 状态、合法边或证明。",
              font=note_font, fill=f"#{DARK_BLUE}")
    image.save(path, dpi=(160, 160))


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, mono: bool = False) -> None:
    name = "Consolas" if mono else "Calibri"
    east_asia = "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start),
                       ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph_border_bottom(paragraph, color: str = BLUE,
                                size: int = 8, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.extend([color, size])
    field_run.append(r_pr)
    field_run.extend([begin, instr, separate, value, end])
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(8)

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Audit Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05

    quote = styles["Quote"]
    quote.font.name = "Calibri"
    quote._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_after = Pt(8)


def add_inline_markup(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, color=DARK_BLUE, mono=True)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Audit Code")
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=8.5, mono=True, color=INK)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [2100, 2700, 4560]
    else:
        widths = [CONTENT_WIDTH_DXA // columns] * columns
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        keep_table_row_together(table.rows[row_index])
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline_markup(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=8.5, bold=row_index == 0,
                             color=DARK_BLUE if row_index == 0 else BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.35))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(9)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_body_from_markdown(doc: Document, text: str,
                           figures: dict[str, tuple[Path, str]]) -> None:
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("[FIGURE:") and stripped.endswith("]"):
            key = stripped[len("[FIGURE:"):-1]
            figure_path, caption = figures[key]
            add_figure(doc, figure_path, caption)
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(3, len(heading_match.group(1)) - 1)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline_markup(paragraph, heading_match.group(2))
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = doc.add_paragraph(style="Quote")
            add_inline_markup(paragraph, stripped[2:])
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markup(paragraph, stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markup(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            following = lines[index].strip()
            if (not following or following.startswith(("#", ">", "-", "|", "```", "[FIGURE:"))
                    or re.match(r"^\d+\.\s+", following)):
                break
            paragraph_lines.append(following)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline_markup(paragraph, " ".join(paragraph_lines))


def build_document() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    lean_figure = ASSETS / "lean-architecture.png"
    runtime_figure = ASSETS / "runtime-pipeline.png"
    build_lean_architecture(lean_figure)
    build_runtime_pipeline(runtime_figure)

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    doc.core_properties.title = "华容道 Lean 4 项目版本架构与检查报告"
    doc.core_properties.subject = "Lean 4 formalization, Mathlib integration, frontend and visualization audit"
    doc.core_properties.author = "HuarongdaoLean project audit"
    doc.core_properties.keywords = "Lean 4, Mathlib, Huarongdao, state space, graph, audit"

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("HUARONGDAOLEAN  |  VERSION AUDIT")
    set_run_font(header_run, size=8.5, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(4)
    kicker_run = kicker.add_run("FORMALIZATION · STATE SPACE · VISUALIZATION")
    set_run_font(kicker_run, size=9, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.add_run("华容道 Lean 4 项目\n版本架构与检查报告")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(
        "从规则语义、Mathlib 图论与群作用，到有限证书、确定性布局和浏览器交互")
    set_run_font(subtitle_run, size=13, color=MUTED)

    metadata = [
        ("审计基线", "main@300fe47"),
        ("审计日期", "2026-08-28"),
        ("Lean / Mathlib", "v4.33.1 / v4.33.1"),
        ("审计目的", "版本检查、证明边界确认、后续研究规划"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    status = doc.add_paragraph()
    status.paragraph_format.space_before = Pt(14)
    status.paragraph_format.space_after = Pt(12)
    status_run = status.add_run(
        "结论：构建与经典证书健康；通用 BFS 可由闭包/距离证书升级，A* 仍只提供经重放验证的可行解。")
    set_run_font(status_run, size=11, bold=True, color=DARK_BLUE)
    set_paragraph_border_bottom(status, BLUE, size=10, space=7)

    overview = doc.add_paragraph()
    overview.paragraph_format.space_before = Pt(8)
    overview_run = overview.add_run(
        "本文档区分四类责任：Lean 定理、Lean 可执行 checker、JavaScript 计算、WebGL 展示。"
        "只有存在明确保持定理或 checker soundness 时，结论才跨越层边界。")
    set_run_font(overview_run, size=11, color=INK)
    doc.add_page_break()

    source_text = SOURCE.read_text(encoding="utf-8")
    figures = {
        "lean-architecture": (lean_figure, "图 1  Lean 语义内核、有限表示、Mathlib 与可执行入口"),
        "runtime-pipeline": (runtime_figure, "图 2  本地求解、布局、渲染与交互的数据链"),
    }
    add_body_from_markdown(doc, source_text, figures)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
